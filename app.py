import streamlit as st
import pandas as pd
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from docx import Document
import io

# =============================
# HEADER APLIKASI
# =============================
st.markdown(
    "<h2 style='text-align:center; color:#4CAF50;'>📂 Kusuma Converter</h2>"
    "<p style='text-align:center;'>Konversi file antara <b>Excel, CSV, Word, dan PDF</b></p>",
    unsafe_allow_html=True
)

# =============================
# SIDEBAR
# =============================
st.sidebar.header("⚙️ Pengaturan")
uploaded_file = st.sidebar.file_uploader("Upload File", type=["xlsx", "csv", "docx"])
output_format = st.sidebar.selectbox("Pilih Format Output", ["CSV", "Excel", "Word", "PDF"])

# =============================
# FUNGSI KONVERSI
# =============================
def df_to_pdf(df):
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    table_data = [df.columns.tolist()] + df.values.tolist()
    table = Table(table_data)

    # Styling tabel PDF
    style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4CAF50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ])
    table.setStyle(style)

    doc.build([table])
    pdf_buffer.seek(0)
    return pdf_buffer

def df_to_word(df):
    buffer = io.BytesIO()
    doc = Document()
    table = doc.add_table(rows=1, cols=len(df.columns))

    # Header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Data
    for row in df.values.tolist():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =============================
# LOGIKA UTAMA
# =============================
if uploaded_file:
    filename = uploaded_file.name

    # Baca file ke DataFrame
    if filename.endswith(".xlsx"):
        df = pd.read_excel(uploaded_file)
    elif filename.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    elif filename.endswith(".docx"):
        doc = Document(uploaded_file)
        data = []
        for table in doc.tables:
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
        df = pd.DataFrame(data[1:], columns=data[0]) if data else pd.DataFrame()
    else:
        st.error("Format file belum didukung")
        df = None

    # Jika DataFrame berhasil dibaca
    if df is not None and not df.empty:
        st.subheader("👀 Preview Data")
        st.dataframe(df.head())

        with st.spinner("⏳ Sedang mengonversi..."):
            if output_format == "CSV":
                csv_data = df.to_csv(index=False).encode("utf-8")
                st.download_button("⬇️ Download CSV", csv_data, "output.csv", "text/csv")
                st.success("Konversi ke CSV selesai ✅")

            elif output_format == "Excel":
                towrite = io.BytesIO()
                df.to_excel(towrite, index=False, engine="openpyxl")
                towrite.seek(0)
                st.download_button("⬇️ Download Excel", towrite, "output.xlsx")
                st.success("Konversi ke Excel selesai ✅")

            elif output_format == "Word":
                word_buffer = df_to_word(df)
                st.download_button("⬇️ Download Word", word_buffer, "output.docx")
                st.success("Konversi ke Word selesai ✅")

            elif output_format == "PDF":
                pdf_buffer = df_to_pdf(df)
                st.download_button("⬇️ Download PDF", pdf_buffer, "output.pdf", "application/pdf")
                st.success("Konversi ke PDF selesai ✅")
    else:
        st.warning("⚠️ Data kosong atau tidak terbaca.")
else:
    st.info("📤 Silakan upload file melalui sidebar untuk mulai konversi.")
